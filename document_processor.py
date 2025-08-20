# ==============================================================================
#                        DOCUMENT PROCESSING UTILITIES
# ==============================================================================
#
# This module provides comprehensive document processing capabilities for the
# Advanced Document Q&A System. It handles downloading, text extraction, and
# intelligent chunking of documents from various formats.
#
# SUPPORTED FORMATS:
# ==================
# - PDF: Using PyMuPDF for efficient text extraction with layout preservation
# - DOCX: Using python-docx for Microsoft Word document processing  
# - EML: Using email library for email message parsing
#
# KEY FEATURES:
# =============
# 1. Multi-format Document Support: Automatic format detection and processing
# 2. Structure-Aware Chunking: Intelligent text splitting that preserves context
# 3. Large Document Handling: Efficient processing of documents >1MB
# 4. Metadata Enrichment: Comprehensive metadata for enhanced retrieval
# 5. Caching System: LRU cache for improved performance on repeated requests
# 6. Error Resilience: Graceful handling of various document processing errors
#
# PROCESSING PIPELINE:
# ====================
# URL → Download → Format Detection → Text Extraction → Chunking → Metadata → Cache
#
# CHUNKING STRATEGY:
# ==================
# The system uses structure-aware chunking that:
# - Detects document headers and sections
# - Preserves semantic boundaries
# - Optimizes chunk sizes for embedding models
# - Maintains context through metadata
# - Handles large documents through sectioning
#
# PERFORMANCE OPTIMIZATIONS:
# ===========================
# - Batched PDF processing to handle large files
# - Memory-efficient text extraction
# - Intelligent section size management
# - LRU caching for document reuse
# - Conservative chunk sizing for optimal embeddings
#
# ==============================================================================

import email
from functools import lru_cache
import os
import docx
import requests
import tempfile
from typing import List, Tuple, Optional
import fitz  # PyMuPDF for PDF processing
import re

from langchain_core.documents import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter

def _extract_text_from_pdf_batched(file_path: str, max_chars_per_batch: int = 500000) -> str:
    """
    Extract text from PDF files using PyMuPDF with intelligent batching.
    
    This function processes large PDF files in manageable batches to avoid
    memory issues while maintaining text extraction quality. It handles
    complex PDF layouts and preserves text structure.
    
    Args:
        file_path (str): Path to the PDF file to process
        max_chars_per_batch (int): Maximum characters per processing batch
                                   (default: 500,000 for optimal memory usage)
    
    Returns:
        str: Complete extracted text from the PDF document
        
    Process:
        1. Open PDF document using PyMuPDF
        2. Calculate optimal batch size based on total pages
        3. Process pages in batches to manage memory
        4. Accumulate text while monitoring batch size
        5. Return complete document text
        
    Note:
        This batched approach is particularly important for large documents
        (>100 pages) where memory usage could become problematic.
    """
    doc = fitz.open(file_path)
    text = ""
    total_pages = len(doc)
    
    # Process pages in batches to avoid memory issues
    # Calculate batch size: aim for ~10 batches for optimal performance
    batch_size = max(1, total_pages // 10)
    current_batch_text = ""
    
    for page_num in range(total_pages):
        page = doc[page_num]
        page_text = page.get_text()
        
        current_batch_text += page_text
        
        # If batch is getting too large, process it and start new batch
        if len(current_batch_text) > max_chars_per_batch:
            text += current_batch_text
            current_batch_text = ""
    
    # Add any remaining text from the last batch
    if current_batch_text:
        text += current_batch_text
    
    doc.close()
    return text

def _extract_text_from_pdf(file_path: str) -> str:
    """Extracts text from a PDF file using PyMuPDF with batching for large documents."""
    try:
        return _extract_text_from_pdf_batched(file_path)
    except Exception as e:
        # Fallback to original method
        doc = fitz.open(file_path)
        text = ""
        for page in doc:
            text += page.get_text()
        doc.close()
        return text



def _extract_text_from_docx(file_path: str) -> str:
    """Extracts text from a DOCX file, including tables."""
    doc = docx.Document(file_path)
    text = []
    for para in doc.paragraphs:
        text.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text.append(cell.text)

    return "\n".join(text)



def _extract_text_from_eml(file_path: str) -> str:
    """Extracts text from an EML file using Python's built-in email library."""
    with open(file_path, 'rb') as f:
        msg = email.message_from_bytes(f.read())

    

    body = ""
    if msg.is_multipart():
        for part in msg.walk():
            content_type = part.get_content_type()

            if content_type == "text/plain":
                payload = part.get_payload(decode=True)
                body += payload.decode(errors='ignore')
    else:
        payload = msg.get_payload(decode=True)
        body = payload.decode(errors='ignore')

    subject = msg.get('subject', '')

    return f"Subject: {subject}\n\n{body}"



def structure_aware_chunking(text: str) -> List[Tuple[str, str]]:
    """
    Splits text into chunks based on section headings (structure-aware).
    Returns a list of (section_header, chunk) tuples.
    Optimized for very large documents and speed.
    """
    # Estimate document size
    text_length = len(text)
    
    # Simplified regex for headings: lines in ALL CAPS or starting with numbers
    heading_regex = re.compile(r'(^[A-Z][A-Z0-9\-\s:]{3,}$|^\d+\. .+)', re.MULTILINE)
    matches = list(heading_regex.finditer(text))
    chunks = []
    
    if not matches:
        # Fallback: split by paragraphs with size limits
        paras = text.split('\n\n')
        current_chunk = ""
        current_size = 0
        max_chunk_size = 1500  # Reduced from 2000 for speed
        
        for para in paras:
            para = para.strip()
            if len(para) > 100:
                if current_size + len(para) > max_chunk_size and current_chunk:
                    chunks.append(("", current_chunk.strip()))
                    current_chunk = para
                    current_size = len(para)
                else:
                    current_chunk += "\n\n" + para if current_chunk else para
                    current_size += len(para)
        
        if current_chunk:
            chunks.append(("", current_chunk.strip()))
        return chunks
    
    # Process by sections
    for i, match in enumerate(matches):
        start = match.end()
        end = matches[i+1].start() if i+1 < len(matches) else len(text)
        section_header = match.group().strip()
        section_text = text[start:end].strip()
        
        # For very large sections, split further
        if len(section_text) > 3000:  # Reduced from 5000
            # Split large sections by paragraphs
            paras = section_text.split('\n\n')
            current_chunk = ""
            current_size = 0
            max_chunk_size = 1500  # Reduced from 2000
            
            for para in paras:
                para = para.strip()
                if len(para) > 100:
                    if current_size + len(para) > max_chunk_size and current_chunk:
                        chunks.append((section_header, current_chunk.strip()))
                        current_chunk = para
                        current_size = len(para)
                    else:
                        current_chunk += "\n\n" + para if current_chunk else para
                        current_size += len(para)
            
            if current_chunk:
                chunks.append((section_header, current_chunk.strip()))
        else:
            # Small section - split by paragraphs
            paras = section_text.split('\n\n')
            for para in paras:
                if len(para.strip()) > 100:
                    chunks.append((section_header, para.strip()))
    
    return chunks


@lru_cache(maxsize=10)
def build_knowledge_base_from_urls(document_url: str) -> List:  
    """
    Downloads a document from URL, extracts text using appropriate parser, and splits
    it into chunks. Supports PDF, Word documents, and email formats.
    Handles very large documents with progressive processing.
    """
    docs_with_metadata: List = []
    
    with tempfile.TemporaryDirectory() as temp_dir:
        try:
            response = requests.get(document_url) 
            response.raise_for_status()
            file_name = os.path.basename(document_url.split('?')[0])  
            if not file_name:
                if '.docx' in document_url.lower():  
                    file_name = "document.docx"
                elif '.eml' in document_url.lower(): 
                    file_name = "document.eml"
                else:
                    file_name = "document.pdf"
            temp_path = os.path.join(temp_dir, file_name)  
            with open(temp_path, 'wb') as f:
                f.write(response.content)
            text_content = ""
            if file_name.lower().endswith('.pdf'):
                text_content = _extract_text_from_pdf(temp_path)
            elif file_name.lower().endswith('.docx'):
                text_content = _extract_text_from_docx(temp_path)
            elif file_name.lower().endswith('.eml'):
                text_content = _extract_text_from_eml(temp_path)
            else:
                return []  
            
            # Check if document is too large and process in chunks if needed
            if len(text_content) > 1000000:  # 1MB threshold
                docs_with_metadata = process_large_document(text_content, document_url, file_name)
            else:
                # Structure-aware chunking for normal documents
                chunk_tuples = structure_aware_chunking(text_content)
                for i, (section_header, chunk) in enumerate(chunk_tuples):
                    docs_with_metadata.append(Document(
                        page_content=chunk,
                        metadata={
                            "source_url": document_url,
                            "section_header": section_header,
                            "chunk_id": i,
                            "file_type": file_name.split('.')[-1].lower() if '.' in file_name else 'unknown'
                        }
                    ))
        except requests.RequestException as e:
            return []  
        except Exception as e:
            return [] 
    if not docs_with_metadata:
        return []  
    return docs_with_metadata

def process_large_document(text_content: str, document_url: str, file_name: str) -> List:
    """
    Process very large documents by splitting them into manageable sections
    before chunking to avoid token limit issues.
    """
    docs_with_metadata = []
    total_length = len(text_content)
    section_size = 500000  # 500KB sections
    section_count = 0
    
    for i in range(0, total_length, section_size):
        section_text = text_content[i:i + section_size]
        section_count += 1
        
        # Process this section with structure-aware chunking
        chunk_tuples = structure_aware_chunking(section_text)
        
        for j, (section_header, chunk) in enumerate(chunk_tuples):
            docs_with_metadata.append(Document(
                page_content=chunk,
                metadata={
                    "source_url": document_url,
                    "section_header": section_header,
                    "chunk_id": f"{section_count}_{j}",
                    "file_type": file_name.split('.')[-1].lower() if '.' in file_name else 'unknown',
                    "section_number": section_count
                }
            ))
    
    return docs_with_metadata